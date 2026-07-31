import io
import json
import shutil
import subprocess
import tarfile
import tempfile
import unittest
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from unittest.mock import patch

from converters.adapters.archive import convert_archive
from converters.adapters.data import dict_to_xml, read_xlsx_rows, xml_to_dict
from converters.adapters.document_basic import _libreoffice_convert
from converters.adapters.image import convert_image
from converters.adapters.media import convert_media
from converters.models import ConversionLevel
from converters.pipeline import convert_file
from converters.executor import convert_file_with_timeout
from converters.registry import find_capability, infer_input_format, list_capabilities


class ConverterRegistryTests(unittest.TestCase):
    def test_registry_contains_large_first_version_capability_matrix(self):
        capabilities = list_capabilities()
        self.assertGreaterEqual(len([c for c in capabilities if c.level == ConversionLevel.STABLE]), 45)
        self.assertGreaterEqual(len([c for c in capabilities if c.level == ConversionLevel.BASIC]), 15)
        self.assertGreaterEqual(len([c for c in capabilities if c.level == ConversionLevel.VENDOR]), 10)

    def test_registry_contains_representative_paths(self):
        self.assertEqual(find_capability('csv', 'xlsx').level, ConversionLevel.STABLE)
        self.assertEqual(find_capability('png', 'webp').level, ConversionLevel.STABLE)
        self.assertEqual(find_capability('mp4', 'mp3').level, ConversionLevel.STABLE)
        self.assertEqual(find_capability('md', 'pdf').level, ConversionLevel.BASIC)
        self.assertEqual(find_capability('pdf', 'docx').level, ConversionLevel.VENDOR)



    def test_registry_expands_non_pdf_document_capabilities(self):
        required = {
            ('txt', 'docx'),
            ('md', 'docx'),
            ('html', 'docx'),
            ('rtf', 'txt'),
            ('rtf', 'html'),
            ('rtf', 'docx'),
            ('docx', 'md'),
            ('docx', 'rtf'),
            ('odt', 'docx'),
            ('doc', 'docx'),
            ('ppt', 'pptx'),
            ('xls', 'xlsx'),
            ('ods', 'xlsx'),
            ('epub', 'txt'),
            ('epub', 'html'),
        }
        for source, target in required:
            with self.subTest(path=f'{source}->{target}'):
                capability = find_capability(source, target)
                self.assertIsNotNone(capability)
                self.assertNotEqual(capability.level, ConversionLevel.VENDOR)
                self.assertNotEqual(capability.target, 'pdf')



    def test_registry_expands_common_non_pdf_media_image_subtitle_paths(self):
        required = {
            ('avi', 'mp4'),
            ('flv', 'mp4'),
            ('m4v', 'mp4'),
            ('3gp', 'mp4'),
            ('ts', 'mp4'),
            ('aiff', 'mp3'),
            ('wma', 'mp3'),
            ('amr', 'mp3'),
            ('png', 'ico'),
            ('ico', 'png'),
            ('ppm', 'png'),
            ('png', 'ppm'),
            ('lrc', 'srt'),
            ('lrc', 'vtt'),
            ('gz', 'folder'),
            ('bz2', 'folder'),
        }
        for source, target in required:
            with self.subTest(path=f'{source}->{target}'):
                capability = find_capability(source, target)
                self.assertIsNotNone(capability)
                self.assertNotEqual(capability.level, ConversionLevel.VENDOR)
                self.assertNotEqual(capability.target, 'pdf')

    def test_infer_input_format_handles_compound_extensions(self):
        self.assertEqual(infer_input_format('archive.tar.gz'), 'tar.gz')
        self.assertEqual(infer_input_format('rows.ndjson'), 'ndjson')
        self.assertEqual(infer_input_format('table.CSV'), 'csv')


class ConverterPipelineTests(unittest.TestCase):
    def test_conversion_timeout_terminates_worker_cleanly(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / 'sample.csv'
            source.write_text('name\nAda\n', encoding='utf-8')
            result = convert_file_with_timeout(
                source,
                source.name,
                'csv',
                'json',
                root,
                timeout_seconds=0,
            )
        self.assertFalse(result.success)
        self.assertIn('超时', result.error)

    def test_pipeline_converts_csv_to_json(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'sample.csv'
            input_file.write_text('name,age\nAda,12\n', encoding='utf-8')

            result = convert_file(input_file, input_file.name, 'csv', 'json', tmp_path)

            self.assertTrue(result.success)
            self.assertIsNotNone(result.output_path)
            self.assertTrue(result.output_path.exists())
            self.assertIn('Ada', result.output_path.read_text(encoding='utf-8'))

    def test_pipeline_converts_json_to_toml_when_tomllib_is_available(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / 'settings.json'
            source.write_text(json.dumps({'app': {'name': 'StreamDock', 'port': 8002}}), encoding='utf-8')

            result = convert_file(source, source.name, 'json', 'toml', root)

            self.assertTrue(result.success, result.error)
            import tomllib
            parsed = tomllib.loads(result.output_path.read_text(encoding='utf-8'))
            self.assertEqual(parsed['app']['name'], 'StreamDock')
            self.assertEqual(parsed['app']['port'], 8002)

    def test_xlsx_reader_preserves_zero_header(self):
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / 'headers.xlsx'
            workbook = Workbook()
            sheet = workbook.active
            sheet.append([0, None])
            sheet.append(['value', 'fallback'])
            workbook.save(path)

            rows = read_xlsx_rows(path)

        self.assertEqual(rows, [{'0': 'value', 'column_2': 'fallback'}])

    def test_xml_reader_preserves_mixed_content_text(self):
        root = ET.fromstring('<root>leading text<child>value</child></root>')

        parsed = xml_to_dict(root)

        self.assertEqual(parsed, {'root': {'#text': 'leading text', 'child': 'value'}})
        rebuilt = dict_to_xml('root', parsed['root'])
        self.assertEqual(rebuilt.text, 'leading text')
        self.assertEqual(rebuilt.findtext('child'), 'value')



    def test_pipeline_converts_txt_to_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'note.txt'
            input_file.write_text('第一行\n第二行\n', encoding='utf-8')

            result = convert_file(input_file, input_file.name, 'txt', 'docx', tmp_path)

            self.assertTrue(result.success, result.error)
            self.assertTrue(result.output_path.exists())
            from docx import Document
            doc = Document(str(result.output_path))
            self.assertIn('第一行', '\n'.join(p.text for p in doc.paragraphs))

    def test_pipeline_converts_markdown_to_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'guide.md'
            input_file.write_text('# 标题\n\n正文内容', encoding='utf-8')

            result = convert_file(input_file, input_file.name, 'md', 'docx', tmp_path)

            self.assertTrue(result.success, result.error)
            self.assertTrue(result.output_path.exists())
            from docx import Document
            doc = Document(str(result.output_path))
            text = '\n'.join(p.text for p in doc.paragraphs)
            self.assertIn('标题', text)
            self.assertIn('正文内容', text)

    def test_pipeline_converts_chinese_markdown_table_to_html(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / '产品思维.md'
            source.write_text('# 产品思维\n\n| 序号 | 项目方向 | 简介 |\n| --- | --- | --- |\n| 1 | 推荐系统 | 中文内容可读 |\n', encoding='utf-8')
            result = convert_file(source, source.name, 'md', 'html', root)
            self.assertTrue(result.success, result.error)
            rendered = result.output_path.read_text(encoding='utf-8')
            self.assertIn('<table>', rendered)
            self.assertIn('中文内容可读', rendered)

    def test_pipeline_converts_chinese_markdown_table_to_readable_pdf(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / '产品思维.md'
            source.write_text('# 产品思维\n\n| 序号 | 项目方向 | 简介 |\n| --- | --- | --- |\n| 1 | 推荐系统 | 中文内容可读 |\n', encoding='utf-8')
            result = convert_file(source, source.name, 'md', 'pdf', root)
            self.assertTrue(result.success, result.error)
            from pypdf import PdfReader
            extracted = '\n'.join(page.extract_text() or '' for page in PdfReader(str(result.output_path)).pages)
            self.assertIn('产品思维', extracted)
            self.assertIn('中文内容可读', extracted)

    def test_pipeline_converts_html_structure_to_markdown(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / 'page.html'
            source.write_text(
                '<h1>Title</h1><p>Read the <a href="https://example.test">guide</a>.</p><ul><li>First</li></ul>',
                encoding='utf-8',
            )

            result = convert_file(source, source.name, 'html', 'md', root)

            self.assertTrue(result.success, result.error)
            markdown = result.output_path.read_text(encoding='utf-8')
            self.assertIn('# Title', markdown)
            self.assertIn('[guide](https://example.test)', markdown)
            self.assertIn('First', markdown)

    def test_pipeline_converts_rtf_to_txt(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'legacy.rtf'
            input_file.write_text(r'{\rtf1\ansi 第一行\par 第二行}', encoding='utf-8')

            result = convert_file(input_file, input_file.name, 'rtf', 'txt', tmp_path)

            self.assertTrue(result.success, result.error)
            content = result.output_path.read_text(encoding='utf-8')
            self.assertIn('第一行', content)
            self.assertIn('第二行', content)

    def test_pipeline_converts_epub_to_html(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'book.epub'
            from ebooklib import epub
            book = epub.EpubBook()
            book.set_identifier('streamdock-test')
            book.set_title('测试书')
            book.set_language('zh')
            chapter = epub.EpubHtml(title='第一章', file_name='chapter.xhtml', lang='zh')
            chapter.content = '<html><body><h1>第一章</h1><p>正文内容</p></body></html>'
            book.add_item(chapter)
            book.toc = (chapter,)
            book.spine = ['nav', chapter]
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            epub.write_epub(str(input_file), book)

            result = convert_file(input_file, input_file.name, 'epub', 'html', tmp_path)

            self.assertTrue(result.success, result.error)
            content = result.output_path.read_text(encoding='utf-8')
            self.assertIn('第一章', content)
            self.assertIn('正文内容', content)



    def test_pipeline_converts_lrc_to_srt(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'song.lrc'
            input_file.write_text('[00:01.20]第一句\n[00:03.00]第二句\n', encoding='utf-8')

            result = convert_file(input_file, input_file.name, 'lrc', 'srt', tmp_path)

            self.assertTrue(result.success, result.error)
            content = result.output_path.read_text(encoding='utf-8')
            self.assertIn('00:00:01,200 --> 00:00:03,000', content)
            self.assertIn('第一句', content)

    def test_pipeline_converts_png_to_ico(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'icon.png'
            from PIL import Image
            Image.new('RGBA', (32, 32), (255, 0, 0, 255)).save(input_file)

            result = convert_file(input_file, input_file.name, 'png', 'ico', tmp_path)

            self.assertTrue(result.success, result.error)
            self.assertTrue(result.output_path.exists())
            self.assertGreater(result.output_path.stat().st_size, 0)

    def test_pipeline_extracts_gz_to_folder(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'hello.txt.gz'
            import gzip
            with gzip.open(input_file, 'wb') as fh:
                fh.write('hello gzip'.encode('utf-8'))

            result = convert_file(input_file, input_file.name, 'gz', 'folder', tmp_path)

            self.assertTrue(result.success, result.error)
            extracted = result.output_path / 'hello.txt'
            self.assertTrue(extracted.exists())
            self.assertEqual(extracted.read_text(encoding='utf-8'), 'hello gzip')

    def test_archive_extraction_rejects_tar_path_traversal(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            archive = tmp_path / 'evil.tar'
            data = b'tar-slip'
            info = tarfile.TarInfo('../pwned.txt')
            info.size = len(data)
            with tarfile.open(archive, 'w') as tf:
                tf.addfile(info, io.BytesIO(data))

            with self.assertRaisesRegex(RuntimeError, '不安全路径'):
                convert_archive('tar', 'folder', archive, tmp_path / 'out')
            self.assertFalse((tmp_path / 'pwned.txt').exists())

    def test_archive_extraction_rejects_zip_path_traversal(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            archive = tmp_path / 'evil.zip'
            with zipfile.ZipFile(archive, 'w') as zf:
                zf.writestr('../pwned.txt', 'zip-slip')

            with self.assertRaisesRegex(RuntimeError, '不安全路径'):
                convert_archive('zip', 'folder', archive, tmp_path / 'out')
            self.assertFalse((tmp_path / 'pwned.txt').exists())

    def test_archive_extraction_rejects_tar_special_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            archive = root / 'special.tar'
            info = tarfile.TarInfo('named-pipe')
            info.type = tarfile.FIFOTYPE
            with tarfile.open(archive, 'w') as tf:
                tf.addfile(info)

            with self.assertRaisesRegex(RuntimeError, '特殊文件'):
                convert_archive('tar', 'folder', archive, root / 'out')

    def test_archive_conversion_removes_intermediate_directory(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            archive = root / 'source.zip'
            output = root / 'converted.tar'
            with zipfile.ZipFile(archive, 'w') as zf:
                zf.writestr('hello.txt', 'hello')

            convert_archive('zip', 'tar', archive, output)

            self.assertTrue(output.is_file())
            self.assertFalse(output.with_suffix('').exists())

    def test_archive_conversion_cleans_intermediate_directory_after_failure(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            archive = root / 'source.tar'
            output = root / 'converted.zip'
            data = b'hello'
            info = tarfile.TarInfo('hello.txt')
            info.size = len(data)
            with tarfile.open(archive, 'w') as tf:
                tf.addfile(info, io.BytesIO(data))

            with patch('converters.adapters.archive._folder_to_zip', side_effect=RuntimeError('zip failed')):
                with self.assertRaisesRegex(RuntimeError, 'zip failed'):
                    convert_archive('tar', 'zip', archive, output)

            self.assertFalse(output.with_suffix('').exists())

    def test_single_file_archive_uses_streaming_copy(self):
        import gzip

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            archive = root / 'large.txt.gz'
            with gzip.open(archive, 'wb') as handle:
                handle.write(b'x' * (1024 * 1024))

            with patch('converters.adapters.archive.shutil.copyfileobj', wraps=shutil.copyfileobj) as copy:
                convert_archive('gz', 'folder', archive, root / 'out')

            copy.assert_called_once()
            self.assertEqual((root / 'out' / 'large.txt').stat().st_size, 1024 * 1024)

    def test_direct_gif_conversion_rejects_silent_first_frame_loss(self):
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / 'animated.gif'
            frames = [Image.new('RGB', (4, 4), color) for color in ('red', 'blue')]
            frames[0].save(source, save_all=True, append_images=frames[1:])

            with self.assertRaisesRegex(RuntimeError, 'PNG 帧序列'):
                convert_image('gif', 'jpg', source, root / 'output.jpg')

    def test_netpbm_targets_use_semantic_pixel_modes(self):
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / 'source.png'
            Image.new('RGB', (3, 3), 'red').save(source)

            pgm = root / 'output.pgm'
            pbm = root / 'output.pbm'
            convert_image('png', 'pgm', source, pgm)
            convert_image('png', 'pbm', source, pbm)

            self.assertEqual(pgm.read_bytes()[:2], b'P5')
            self.assertEqual(pbm.read_bytes()[:2], b'P4')

    def test_pipeline_rejects_vendor_only_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'sample.pdf'
            input_file.write_bytes(b'%PDF-1.4')

            result = convert_file(input_file, input_file.name, 'pdf', 'docx', tmp_path)

            self.assertFalse(result.success)
            self.assertTrue(result.vendor_recommendations)
            self.assertIn('推荐厂商', result.error)

    def test_media_conversion_passes_timeout_to_ffmpeg(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'demo.mp4'
            output_file = tmp_path / 'demo.mp3'
            input_file.write_bytes(b'fake')

            with patch('converters.adapters.media.shutil.which', return_value='/usr/bin/ffmpeg'):
                with patch('converters.adapters.media.subprocess.run') as mocked_run:
                    mocked_run.return_value.returncode = 0
                    mocked_run.return_value.stderr = ''

                    convert_media('mp4', 'mp3', input_file, output_file)

            self.assertIn('timeout', mocked_run.call_args.kwargs)
            self.assertGreater(mocked_run.call_args.kwargs['timeout'], 0)

    def test_liberoffice_conversion_reports_timeout_cleanly(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            input_file = tmp_path / 'demo.docx'
            input_file.write_bytes(b'fake')

            with patch('converters.adapters.document_basic.shutil.which', return_value='/usr/bin/soffice'):
                with patch(
                    'converters.adapters.document_basic.subprocess.run',
                    side_effect=subprocess.TimeoutExpired(cmd='soffice', timeout=1),
                ):
                    with self.assertRaisesRegex(RuntimeError, '超时'):
                        _libreoffice_convert(input_file, tmp_path, 'pdf')


if __name__ == '__main__':
    unittest.main()
