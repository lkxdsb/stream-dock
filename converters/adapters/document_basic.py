from __future__ import annotations

import html
import os
import re
import shutil
import subprocess
from pathlib import Path

LIBREOFFICE_CONVERT_TIMEOUT_SECONDS = int(os.getenv('STREAMDOCK_LIBREOFFICE_TIMEOUT_SECONDS', str(10 * 60)))


def _html_to_text(text: str) -> str:
    text = re.sub(r'<\s*br\s*/?>', '\n', text, flags=re.I)
    text = re.sub(r'</\s*p\s*>', '\n\n', text, flags=re.I)
    text = re.sub(r'<[^>]+>', '', text)
    return html.unescape(text).strip() + '\n'


def _html_to_markdown(text: str) -> str:
    try:
        from markdownify import markdownify  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError('缺少 markdownify，无法将 HTML 转换为 Markdown') from exc
    return markdownify(text, heading_style='ATX').strip() + '\n'


def _markdown_to_text(text: str) -> str:
    text = re.sub(r'```.*?```', '', text, flags=re.S)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'^\s{0,3}#{1,6}\s*', '', text, flags=re.M)
    text = re.sub(r'[*_>~]', '', text)
    text = re.sub(r'^\s*[-+]\s+', '- ', text, flags=re.M)
    return text.strip() + '\n'


def _rtf_to_text(text: str) -> str:
    text = re.sub(r'\\par[d]?', '\n', text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", '', text)
    text = re.sub(r'\\u(-?\d+)\??', lambda m: chr(int(m.group(1)) % 65536), text)
    text = re.sub(r'\\[a-zA-Z]+-?\d* ?', '', text)
    text = text.replace('{', '').replace('}', '')
    text = text.replace('\\~', ' ').replace('\\-', '')
    lines = [line.strip() for line in html.unescape(text).splitlines()]
    return '\n'.join(line for line in lines if line).strip() + '\n'


def _text_to_html(text: str) -> str:
    body = ''.join(f'<p>{html.escape(line)}</p>\n' for line in text.splitlines() if line.strip())
    return '<!doctype html><meta charset="utf-8">\n' + body


def _text_to_rtf(text: str) -> str:
    escaped = text.replace('\\', r'\\').replace('{', r'\{').replace('}', r'\}')
    escaped = escaped.replace('\n', r'\par ' + '\n')
    return r'{\rtf1\ansi\deff0 ' + escaped + '}'


def _write_docx_from_text(text: str, output_path: Path) -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError('缺少 python-docx，无法生成 DOCX') from exc
    doc = Document()
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph('')
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)
    doc.save(str(output_path))


def _read_docx_paragraphs(input_path: Path) -> list[str]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError('缺少 python-docx，无法读取 DOCX') from exc
    doc = Document(str(input_path))
    return [p.text for p in doc.paragraphs]


def _markdown_to_html(text: str) -> str:
    try:
        import markdown  # type: ignore
    except ImportError:
        markdown = None
    if markdown is not None:
        body = markdown.markdown(text, extensions=['tables', 'fenced_code'])
    else:
        lines = []
        for line in text.splitlines():
            if line.startswith('# '):
                lines.append(f'<h1>{html.escape(line[2:])}</h1>')
            elif line.startswith('## '):
                lines.append(f'<h2>{html.escape(line[3:])}</h2>')
            elif line.strip():
                lines.append(f'<p>{html.escape(line)}</p>')
        body = '\n'.join(lines)
    return '<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><title>StreamDock 文档</title></head><body>\n' + body + '\n</body></html>\n'


def _write_basic_pdf(text: str, output_path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError('缺少 reportlab，无法生成基础 PDF') from exc
    c = canvas.Canvas(str(output_path), pagesize=A4)
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    c.setFont('STSong-Light', 10)
    width, height = A4
    y = height - 50
    for raw_line in text.splitlines():
        line = raw_line or ' '
        chunks = [line[index:index + 55] for index in range(0, len(line), 55)] or [' ']
        for chunk in chunks:
            if y < 50:
                c.showPage(); c.setFont('STSong-Light', 10); y = height - 50
            c.drawString(50, y, chunk)
            y -= 16
    c.save()


def _write_markdown_pdf(text: str, output_path: Path) -> None:
    """Render readable Chinese Markdown, including wide pipe tables, to PDF."""
    try:
        from reportlab.lib import colors  # type: ignore
        from reportlab.lib.pagesizes import A4, landscape  # type: ignore
        from reportlab.lib.styles import ParagraphStyle  # type: ignore
        from reportlab.lib.enums import TA_LEFT  # type: ignore
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
        from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, TableStyle  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError('缺少 reportlab，无法生成 Markdown PDF') from exc

    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    page_size = landscape(A4)
    document = SimpleDocTemplate(str(output_path), pagesize=page_size, leftMargin=30, rightMargin=30, topMargin=32, bottomMargin=32)
    body = ParagraphStyle('CJKBody', fontName='STSong-Light', fontSize=9, leading=14, alignment=TA_LEFT, wordWrap='CJK')
    heading = ParagraphStyle('CJKHeading', parent=body, fontSize=20, leading=26, spaceAfter=10)
    subheading = ParagraphStyle('CJKSubheading', parent=body, fontSize=14, leading=20, spaceBefore=8, spaceAfter=6)
    table_text = ParagraphStyle('CJKTable', parent=body, fontSize=6.4, leading=9, wordWrap='CJK')
    story = []
    lines = text.splitlines()
    index = 0

    def cells(line: str) -> list[str]:
        return [part.strip() for part in line.strip().strip('|').split('|')]

    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith('|') and index + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-{3,}', lines[index + 1]):
            rows = [cells(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith('|'):
                rows.append(cells(lines[index])); index += 1
            column_count = max(len(row) for row in rows)
            normalized = [row + [''] * (column_count - len(row)) for row in rows]
            available_width = page_size[0] - 60
            table = LongTable(
                [[Paragraph(html.escape(value), table_text) for value in row] for row in normalized],
                colWidths=[available_width / column_count] * column_count,
                repeatRows=1,
            )
            table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'STSong-Light'), ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eee8df')),
                ('GRID', (0, 0), (-1, -1), .35, colors.HexColor('#cfc7bd')), ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 4), ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 4), ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            story.extend([table, Spacer(1, 10)])
            continue
        if line.startswith('# '):
            story.append(Paragraph(html.escape(line[2:]), heading))
        elif line.startswith(('## ', '### ')):
            story.append(Paragraph(html.escape(line.lstrip('#').strip()), subheading))
        else:
            story.extend([Paragraph(html.escape(line), body), Spacer(1, 5)])
        index += 1
    document.build(story)


def _libreoffice_convert(input_path: Path, output_dir: Path, target: str) -> Path:
    bundled_soffice = Path.home() / '.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/soffice'
    binary = shutil.which('soffice') or shutil.which('libreoffice') or (str(bundled_soffice) if bundled_soffice.exists() else None)
    if not binary:
        raise RuntimeError('缺少 LibreOffice，无法执行 Office 基础转换')
    try:
        completed = subprocess.run(
            [binary, '--headless', '--convert-to', target, '--outdir', str(output_dir), str(input_path)],
            text=True,
            capture_output=True,
            timeout=LIBREOFFICE_CONVERT_TIMEOUT_SECONDS,
        )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError(f'LibreOffice 转换超时，已停止任务（{LIBREOFFICE_CONVERT_TIMEOUT_SECONDS} 秒）') from exc
    if completed.returncode != 0:
        raise RuntimeError(f'LibreOffice 转换失败：{completed.stderr.strip() or completed.stdout.strip()}')
    produced = output_dir / f'{input_path.stem}.{target.split(":", 1)[0]}'
    if not produced.exists():
        matches = list(output_dir.glob(f'{input_path.stem}.*'))
        if matches:
            return matches[0]
        raise RuntimeError('LibreOffice 未生成输出文件')
    return produced


def _convert_with_libreoffice(input_path: Path, output_path: Path, target: str) -> None:
    produced = _libreoffice_convert(input_path, output_path.parent, target)
    if produced != output_path and produced.exists():
        produced.replace(output_path)


def _epub_documents(input_path: Path) -> list[str]:
    try:
        import ebooklib  # type: ignore
        from ebooklib import epub  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError('缺少 ebooklib，无法读取 EPUB') from exc
    book = epub.read_epub(str(input_path))
    documents = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            documents.append(item.get_content().decode('utf-8', errors='ignore'))
    return documents


def convert_document_basic(source: str, target: str, input_path: Path, output_path: Path) -> list[str]:
    text = input_path.read_text(encoding='utf-8', errors='ignore') if source in {'txt', 'md', 'markdown', 'html', 'rtf'} else ''
    if source in {'md', 'markdown'} and target == 'html':
        output_path.write_text(_markdown_to_html(text), encoding='utf-8')
    elif source in {'md', 'markdown'} and target == 'txt':
        output_path.write_text(_markdown_to_text(text), encoding='utf-8')
    elif source in {'md', 'markdown'} and target == 'docx':
        _write_docx_from_text(text, output_path)
    elif source in {'md', 'markdown'} and target == 'pdf':
        _write_markdown_pdf(text, output_path)
    elif source == 'html' and target == 'txt':
        output_path.write_text(_html_to_text(text), encoding='utf-8')
    elif source == 'html' and target == 'md':
        output_path.write_text(_html_to_markdown(text), encoding='utf-8')
    elif source == 'html' and target == 'docx':
        _write_docx_from_text(_html_to_text(text), output_path)
    elif source == 'html' and target == 'pdf':
        _write_basic_pdf(_html_to_text(text), output_path)
    elif source == 'txt' and target == 'html':
        output_path.write_text(_text_to_html(text), encoding='utf-8')
    elif source == 'txt' and target == 'md':
        output_path.write_text(text, encoding='utf-8')
    elif source == 'txt' and target == 'docx':
        _write_docx_from_text(text, output_path)
    elif source == 'txt' and target == 'rtf':
        output_path.write_text(_text_to_rtf(text), encoding='utf-8')
    elif source == 'txt' and target == 'pdf':
        _write_basic_pdf(text, output_path)
    elif source == 'rtf' and target in {'txt', 'html', 'docx'}:
        plain = _rtf_to_text(text)
        if target == 'txt':
            output_path.write_text(plain, encoding='utf-8')
        elif target == 'html':
            output_path.write_text(_text_to_html(plain), encoding='utf-8')
        else:
            _write_docx_from_text(plain, output_path)
    elif source == 'docx' and target in {'txt', 'html'}:
        paragraphs = _read_docx_paragraphs(input_path)
        if target == 'txt':
            output_path.write_text('\n'.join(paragraphs), encoding='utf-8')
        else:
            output_path.write_text('<!doctype html><meta charset="utf-8">\n' + ''.join(f'<p>{html.escape(p)}</p>' for p in paragraphs), encoding='utf-8')
    elif source == 'docx' and target == 'md':
        paragraphs = _read_docx_paragraphs(input_path)
        output_path.write_text('\n\n'.join(p for p in paragraphs if p.strip()) + '\n', encoding='utf-8')
    elif source == 'docx' and target == 'rtf':
        paragraphs = _read_docx_paragraphs(input_path)
        output_path.write_text(_text_to_rtf('\n'.join(paragraphs)), encoding='utf-8')
    elif source in {'docx', 'pptx', 'xlsx'} and target in {'pdf', 'html', 'png'}:
        _convert_with_libreoffice(input_path, output_path, 'pdf' if target == 'png' else target)
    elif source in {'doc', 'odt'} and target in {'docx', 'txt', 'html'}:
        _convert_with_libreoffice(input_path, output_path, target)
    elif source in {'ppt', 'odp'} and target == 'pptx':
        _convert_with_libreoffice(input_path, output_path, target)
    elif source in {'xls', 'ods'} and target in {'xlsx', 'csv'}:
        _convert_with_libreoffice(input_path, output_path, target)
    elif source == 'svg' and target in {'png', 'jpg', 'pdf'}:
        try:
            import cairosvg  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError('缺少 cairosvg，无法转换 SVG') from exc
        if target == 'png':
            cairosvg.svg2png(url=str(input_path), write_to=str(output_path))
        elif target == 'pdf':
            cairosvg.svg2pdf(url=str(input_path), write_to=str(output_path))
        else:
            png_path = output_path.with_suffix('.png')
            cairosvg.svg2png(url=str(input_path), write_to=str(png_path))
            from .image import convert_image
            convert_image('png', 'jpg', png_path, output_path)
    elif source == 'epub' and target in {'txt', 'html', 'md', 'pdf'}:
        docs = _epub_documents(input_path)
        joined_html = '<!doctype html><meta charset="utf-8">\n' + '\n<hr />\n'.join(docs)
        plain = _html_to_text(joined_html)
        if target == 'html':
            output_path.write_text(joined_html, encoding='utf-8')
        elif target == 'txt':
            output_path.write_text(plain, encoding='utf-8')
        elif target == 'md':
            output_path.write_text(_html_to_markdown(joined_html), encoding='utf-8')
        else:
            _write_basic_pdf(plain, output_path)
    else:
        raise RuntimeError(f'暂不支持基础文档转换 {source} → {target}')
    return [f'基础文档转换完成：{source.upper()} → {target.upper()}']
